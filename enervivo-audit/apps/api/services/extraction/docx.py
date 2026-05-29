"""Extraction DOCX/DOTX via python-docx (les .dotx sont des templates OOXML
   au même format que .docx, python-docx les ouvre nativement)."""

from __future__ import annotations

import io

from docx import Document

from .base import ExtractionError, TextExtractor


class DocxExtractor(TextExtractor):
    def extract(self, content: bytes) -> str:
        # Mock : texte brut
        if not content.startswith(b"PK"):
            return content.decode("utf-8", errors="ignore")
        try:
            doc = Document(io.BytesIO(content))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    paras.extend(cell.text for cell in row.cells if cell.text.strip())
            text = "\n".join(paras)
        except Exception as e:
            raise ExtractionError(f"python-docx error : {e}") from e
        if not text.strip():
            raise ExtractionError("DOCX vide")
        return text
